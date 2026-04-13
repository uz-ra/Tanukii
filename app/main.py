import os
import tempfile
import json
import time
import shutil
import threading
import subprocess
import base64
from pathlib import Path
from typing import Optional
from uuid import uuid4
import requests

from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from fastapi.staticfiles import StaticFiles
from faster_whisper import WhisperModel
from pydantic import BaseModel

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

try:
    from docx import Document
except Exception:
    Document = None

try:
    from openai import OpenAI
except Exception:
    OpenAI = None

Llama = None
hf_hub_download = None


BASE_DIR = Path(__file__).resolve().parent.parent
WEB_DIR = BASE_DIR / "web"
MODEL_DIR = BASE_DIR / "models"
MODEL_DIR.mkdir(parents=True, exist_ok=True)
CONFIG_DIR = BASE_DIR / "config"
CONFIG_PATH = CONFIG_DIR / "settings.json"

WHISPER_MODEL_NAME = os.getenv("WHISPER_MODEL", "small")
WHISPER_COMPUTE_TYPE = os.getenv("WHISPER_COMPUTE_TYPE", "int8")
AVAILABLE_WHISPER_MODELS = ["tiny", "base", "small", "medium", "large"]
SUMMARY_PROVIDER = os.getenv("SUMMARY_PROVIDER", "auto").lower()
OPENAI_MODEL = os.getenv("OPENAI_SUMMARY_MODEL", "gpt-4.1-mini")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
GEMINI_MODEL = os.getenv("GEMINI_SUMMARY_MODEL", "gemini-2.0-flash")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
LOCAL_SUMMARY_MODEL = os.getenv("LOCAL_SUMMARY_MODEL", "alfredplpl/gemma-2-2b-jpn-it-gguf")
LOCAL_SUMMARY_GGUF_FILE = os.getenv("LOCAL_SUMMARY_GGUF_FILE", "gemma-2-2b-jpn-it-Q4_K_M.gguf")
LOCAL_SUMMARY_MAX_NEW_TOKENS = int(os.getenv("LOCAL_SUMMARY_MAX_NEW_TOKENS", "384"))
LOCAL_SUMMARY_TEMPERATURE = float(os.getenv("LOCAL_SUMMARY_TEMPERATURE", "0.2"))
LOCAL_SUMMARY_CONTEXT_LENGTH = int(os.getenv("LOCAL_SUMMARY_CONTEXT_LENGTH", "4096"))
LOCAL_SUMMARY_THREADS = int(os.getenv("LOCAL_SUMMARY_THREADS", str(max(1, (os.cpu_count() or 4) // 2))))
DEFAULT_SUMMARY_SYSTEM_PROMPT = (
    "あなたは日本語の会議要約アシスタントです。"
    "冗長さを避け、実務向けにまとめてください。"
    "Markdown記法を使わず、プレーンテキストで出力してください。"
    "出力は要約のみで、指示に対する返答などは含まないでください。"
)
DEFAULT_SUMMARY_USER_PROMPT_TEMPLATE = "{style_instruction}\n\n対象テキスト:\n{text}"

DEFAULT_CONFIG = {
    "whisper_model": WHISPER_MODEL_NAME,
    "debug_mode": False,
    "summary_system_prompt": DEFAULT_SUMMARY_SYSTEM_PROMPT,
    "summary_user_prompt_template": DEFAULT_SUMMARY_USER_PROMPT_TEMPLATE,
    "summary_provider": SUMMARY_PROVIDER,
    "openai_model": OPENAI_MODEL,
    "openai_api_key": OPENAI_API_KEY or "",
    "gemini_model": GEMINI_MODEL,
    "gemini_api_key": GEMINI_API_KEY or "",
    "local_summary_model": LOCAL_SUMMARY_MODEL,
    "local_summary_gguf_file": LOCAL_SUMMARY_GGUF_FILE,
    "local_summary_max_new_tokens": LOCAL_SUMMARY_MAX_NEW_TOKENS,
    "local_summary_temperature": LOCAL_SUMMARY_TEMPERATURE,
    "local_summary_context_length": LOCAL_SUMMARY_CONTEXT_LENGTH,
    "local_summary_threads": LOCAL_SUMMARY_THREADS,
}


app = FastAPI(title="Local Transcribe + Summarize", version="0.1.0")

cors_origins_raw = os.getenv("CORS_ORIGINS", "*")
cors_origins = [origin.strip() for origin in cors_origins_raw.split(",") if origin.strip()]
app.add_middleware(
    CORSMiddleware,
    allow_origins=cors_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


_whisper_models: dict[str, WhisperModel] = {}
_debug_logs: list[dict] = []
_DEBUG_LOG_LIMIT = 300
_transcribe_jobs: dict[str, dict] = {}
_transcribe_jobs_lock = threading.Lock()
_TRANSCRIBE_JOB_LIMIT = 100
_TRANSCRIBE_CHUNK_SECONDS = 300
_TRANSCRIBE_CHUNK_RETRIES = 3
_TRANSCRIBE_PROGRESS_STEP_SECONDS = 5.0
_local_summary_model_cache: dict[str, dict] = {}
_local_summary_model_lock = threading.Lock()


def add_debug_log(level: str, message: str) -> None:
    entry = {
        "ts": time.strftime("%Y-%m-%d %H:%M:%S"),
        "level": level,
        "message": message,
    }
    _debug_logs.append(entry)
    if len(_debug_logs) > _DEBUG_LOG_LIMIT:
        del _debug_logs[: len(_debug_logs) - _DEBUG_LOG_LIMIT]


def update_transcribe_job(job_id: str, **fields) -> None:
    with _transcribe_jobs_lock:
        job = _transcribe_jobs.get(job_id)
        if not job:
            return
        job.update(fields)
        job["updated_at"] = time.time()


def create_transcribe_job(selected_model: str, language: str, filename: str, initial_prompt: str = "") -> str:
    job_id = str(uuid4())
    now = time.time()
    with _transcribe_jobs_lock:
        if len(_transcribe_jobs) >= _TRANSCRIBE_JOB_LIMIT:
            oldest = sorted(_transcribe_jobs.items(), key=lambda item: item[1].get("updated_at", 0.0))
            for old_job_id, _ in oldest[: max(1, len(_transcribe_jobs) - _TRANSCRIBE_JOB_LIMIT + 1)]:
                _transcribe_jobs.pop(old_job_id, None)

        _transcribe_jobs[job_id] = {
            "job_id": job_id,
            "status": "queued",
            "progress": 0,
            "message": "キューに登録しました。",
            "model": selected_model,
            "language": language,
            "filename": filename,
            "initial_prompt": initial_prompt,
            "result": None,
            "error": None,
            "created_at": now,
            "updated_at": now,
        }
    return job_id


def get_transcribe_job(job_id: str) -> Optional[dict]:
    with _transcribe_jobs_lock:
        job = _transcribe_jobs.get(job_id)
        if not job:
            return None
        return dict(job)


def run_transcribe_job(job_id: str, tmp_path: str, selected_model: str, language: str, initial_prompt: str = "") -> None:
    update_transcribe_job(job_id, status="running", message="モデルを読み込み中です。", progress=1)

    try:
        add_debug_log("debug", f"job={job_id} loading whisper model")
        probed_duration = probe_audio_duration_seconds(tmp_path)
        if probed_duration:
            planned_chunks = max(1, int((probed_duration + _TRANSCRIBE_CHUNK_SECONDS - 1) // _TRANSCRIBE_CHUNK_SECONDS))
            if planned_chunks > 1:
                update_transcribe_job(
                    job_id,
                    message=(
                        "音声ファイルが大きいので分割モードで実行します。"
                        f" チャンク数={planned_chunks}"
                    ),
                    progress=2,
                )
            else:
                update_transcribe_job(
                    job_id,
                    message="分割モードは使用しません。理由: 音声長が5分以下。",
                    progress=2,
                )
        else:
            update_transcribe_job(
                job_id,
                message="分割判定に必要な音声長を取得できませんでした。フォールバックで処理します。",
                progress=2,
            )

        transcribed = transcribe_with_chunk_restart(
            input_path=tmp_path,
            selected_model=selected_model,
            language=language,
            initial_prompt=initial_prompt,
            progress_cb=lambda progress, chunk_no, chunk_total, end_sec: update_transcribe_job(
                job_id,
                progress=max(2, progress),
                message=(
                    f"文字起こし中... {max(2, progress)}% "
                    f"(チャンク {chunk_no}/{chunk_total}, {round(end_sec, 1)}秒まで完了)"
                ),
            ),
            logger=lambda level, msg: add_debug_log(level, f"job={job_id} {msg}"),
        )

        result = {
            "model": selected_model,
            "language": transcribed.get("language", language),
            "duration": transcribed.get("duration"),
            "text": transcribed.get("text", ""),
            "segments": transcribed.get("segments", []),
        }

        update_transcribe_job(
            job_id,
            status="completed",
            progress=100,
            message="文字起こしが完了しました。",
            result=result,
            error=None,
        )
        add_debug_log(
            "info",
            (
                f"transcribe done: job={job_id}, language={transcribed.get('language', language)}, "
                f"segments={len(transcribed.get('segments', []))}, chars={len(transcribed.get('text', ''))}"
            ),
        )
    except Exception as exc:
        update_transcribe_job(
            job_id,
            status="failed",
            message="文字起こしに失敗しました。",
            error=str(exc),
        )
        add_debug_log("error", f"transcribe failed: job={job_id}, error={str(exc)}")
    finally:
        try:
            os.remove(tmp_path)
            add_debug_log("debug", f"temporary file removed: {tmp_path}")
        except OSError:
            pass


class AppConfigPayload(BaseModel):
    whisper_model: str = WHISPER_MODEL_NAME
    debug_mode: bool = False
    summary_system_prompt: str = DEFAULT_SUMMARY_SYSTEM_PROMPT
    summary_user_prompt_template: str = DEFAULT_SUMMARY_USER_PROMPT_TEMPLATE
    summary_provider: str = "auto"
    openai_model: str = OPENAI_MODEL
    openai_api_key: str = ""
    gemini_model: str = GEMINI_MODEL
    gemini_api_key: str = ""
    local_summary_model: str = LOCAL_SUMMARY_MODEL
    local_summary_gguf_file: str = LOCAL_SUMMARY_GGUF_FILE
    local_summary_max_new_tokens: int = LOCAL_SUMMARY_MAX_NEW_TOKENS
    local_summary_temperature: float = LOCAL_SUMMARY_TEMPERATURE
    local_summary_context_length: int = LOCAL_SUMMARY_CONTEXT_LENGTH
    local_summary_threads: int = LOCAL_SUMMARY_THREADS


def read_config() -> dict:
    if not CONFIG_PATH.exists():
        return dict(DEFAULT_CONFIG)

    try:
        with CONFIG_PATH.open("r", encoding="utf-8") as f:
            loaded = json.load(f)
            if not isinstance(loaded, dict):
                return dict(DEFAULT_CONFIG)
    except Exception:
        return dict(DEFAULT_CONFIG)

    merged = dict(DEFAULT_CONFIG)
    string_keys = {
        "whisper_model",
        "summary_system_prompt",
        "summary_user_prompt_template",
        "summary_provider",
        "openai_model",
        "openai_api_key",
        "gemini_model",
        "gemini_api_key",
        "local_summary_model",
        "local_summary_gguf_file",
    }
    for key in string_keys:
        if key in loaded and isinstance(loaded[key], str):
            merged[key] = loaded[key].strip()

    if "debug_mode" in loaded:
        merged["debug_mode"] = bool(loaded["debug_mode"])

    if "local_summary_max_new_tokens" in loaded:
        try:
            merged["local_summary_max_new_tokens"] = max(64, int(loaded["local_summary_max_new_tokens"]))
        except Exception:
            pass

    if "local_summary_context_length" in loaded:
        try:
            merged["local_summary_context_length"] = max(2048, int(loaded["local_summary_context_length"]))
        except Exception:
            pass

    if "local_summary_threads" in loaded:
        try:
            merged["local_summary_threads"] = max(1, int(loaded["local_summary_threads"]))
        except Exception:
            pass

    if "local_summary_temperature" in loaded:
        try:
            temp = float(loaded["local_summary_temperature"])
            merged["local_summary_temperature"] = min(1.5, max(0.0, temp))
        except Exception:
            pass

    # Migrate legacy default that pointed to RakutenAI transformers model.
    if merged.get("local_summary_model") == "Rakuten/RakutenAI-7B-Instruct":
        merged["local_summary_model"] = LOCAL_SUMMARY_MODEL

    if not merged.get("local_summary_gguf_file"):
        merged["local_summary_gguf_file"] = LOCAL_SUMMARY_GGUF_FILE

    return merged


def write_config(payload: AppConfigPayload) -> dict:
    CONFIG_DIR.mkdir(parents=True, exist_ok=True)
    data = {
        "whisper_model": payload.whisper_model.strip().lower() or WHISPER_MODEL_NAME,
        "debug_mode": bool(payload.debug_mode),
        "summary_system_prompt": payload.summary_system_prompt.strip() or DEFAULT_SUMMARY_SYSTEM_PROMPT,
        "summary_user_prompt_template": payload.summary_user_prompt_template.strip() or DEFAULT_SUMMARY_USER_PROMPT_TEMPLATE,
        "summary_provider": payload.summary_provider.strip().lower() or "auto",
        "openai_model": payload.openai_model.strip() or OPENAI_MODEL,
        "openai_api_key": payload.openai_api_key.strip(),
        "gemini_model": payload.gemini_model.strip() or GEMINI_MODEL,
        "gemini_api_key": payload.gemini_api_key.strip(),
        "local_summary_model": payload.local_summary_model.strip() or LOCAL_SUMMARY_MODEL,
        "local_summary_gguf_file": payload.local_summary_gguf_file.strip() or LOCAL_SUMMARY_GGUF_FILE,
        "local_summary_max_new_tokens": max(64, int(payload.local_summary_max_new_tokens)),
        "local_summary_temperature": min(1.5, max(0.0, float(payload.local_summary_temperature))),
        "local_summary_context_length": max(2048, int(payload.local_summary_context_length)),
        "local_summary_threads": max(1, int(payload.local_summary_threads)),
    }
    if data["whisper_model"] not in AVAILABLE_WHISPER_MODELS:
        data["whisper_model"] = WHISPER_MODEL_NAME
    if data["summary_provider"] not in {"auto", "openai", "gemini", "local"}:
        data["summary_provider"] = "auto"

    with CONFIG_PATH.open("w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    return data


def ffmpeg_available() -> bool:
    return shutil.which("ffmpeg") is not None


def get_whisper_model(model_name: str) -> WhisperModel:
    if model_name not in AVAILABLE_WHISPER_MODELS:
        raise ValueError(f"対応外モデルです: {model_name}")
    if model_name not in _whisper_models:
        _whisper_models[model_name] = WhisperModel(
            model_name,
            download_root=str(MODEL_DIR),
            compute_type=WHISPER_COMPUTE_TYPE,
        )
    return _whisper_models[model_name]


def probe_audio_duration_seconds(file_path: str) -> Optional[float]:
    cmd = [
        "ffprobe",
        "-v",
        "error",
        "-show_entries",
        "format=duration",
        "-of",
        "default=noprint_wrappers=1:nokey=1",
        file_path,
    ]
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, check=True)
        value = (result.stdout or "").strip()
        if not value:
            return None
        seconds = float(value)
        if seconds <= 0:
            return None
        return seconds
    except Exception:
        return None


def extract_audio_chunk(source_path: str, start_sec: float, duration_sec: float, suffix: str) -> str:
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as chunk_tmp:
        chunk_path = chunk_tmp.name

    cmd = [
        "ffmpeg",
        "-hide_banner",
        "-loglevel",
        "error",
        "-y",
        "-ss",
        str(start_sec),
        "-i",
        source_path,
        "-t",
        str(duration_sec),
        "-ac",
        "1",
        "-ar",
        "16000",
        chunk_path,
    ]
    try:
        subprocess.run(cmd, capture_output=True, text=True, check=True)
        return chunk_path
    except Exception:
        try:
            os.remove(chunk_path)
        except OSError:
            pass
        raise


def transcribe_with_chunk_restart(
    input_path: str,
    selected_model: str,
    language: str,
    initial_prompt: str = "",
    progress_cb=None,
    logger=None,
) -> dict:
    if logger:
        logger(
            "info",
            (
                f"chunked transcribe init: model={selected_model}, language={language}, "
                f"chunk_seconds={_TRANSCRIBE_CHUNK_SECONDS}, retries={_TRANSCRIBE_CHUNK_RETRIES}"
            ),
        )

    whisper_model = get_whisper_model(selected_model)
    suffix = Path(input_path).suffix or ".tmp"

    total_duration = probe_audio_duration_seconds(input_path)
    duration_probe_ok = bool(total_duration)
    if not total_duration:
        if logger:
            logger(
                "warning",
                (
                    "audio duration probe failed; "
                    f"fallback duration={_TRANSCRIBE_CHUNK_SECONDS}s"
                ),
            )
        total_duration = float(_TRANSCRIBE_CHUNK_SECONDS)
    elif logger:
        logger("debug", f"audio duration detected: {round(total_duration, 2)}s")

    chunk_size = float(_TRANSCRIBE_CHUNK_SECONDS)
    total_chunks = max(1, int((total_duration + chunk_size - 1) // chunk_size))
    if logger:
        logger(
            "info",
            (
                f"chunk plan: total_duration={round(total_duration, 2)}s, "
                f"chunk_size={int(chunk_size)}s, chunks={total_chunks}"
            ),
        )
        if total_chunks > 1:
            logger(
                "info",
                (
                    "音声ファイルが大きいので分割モードで実行します。"
                    f" チャンク数={total_chunks} (1チャンク{int(chunk_size)}秒)"
                ),
            )
        else:
            reason = (
                "音声長が5分以下"
                if duration_probe_ok
                else "音声長を取得できず分割判定不可(フォールバック)"
            )
            logger(
                "info",
                f"分割モードは使用しません。理由: {reason}。",
            )

    transcript_parts = []
    timed_segments = []
    detected_language: Optional[str] = None

    for chunk_index in range(total_chunks):
        start_sec = float(chunk_index * _TRANSCRIBE_CHUNK_SECONDS)
        remaining = max(0.0, total_duration - start_sec)
        duration_sec = chunk_size if remaining <= 0 else min(chunk_size, remaining)
        if duration_sec <= 0:
            break

        if logger:
            logger(
                "debug",
                f"transcribe chunk start: chunk={chunk_index + 1}/{total_chunks}, start={round(start_sec, 2)}s, duration={round(duration_sec, 2)}s",
            )

        last_error = None
        for attempt in range(1, _TRANSCRIBE_CHUNK_RETRIES + 1):
            chunk_path = ""
            try:
                if logger and attempt > 1:
                    logger(
                        "warning",
                        (
                            f"chunk retry: chunk={chunk_index + 1}/{total_chunks}, "
                            f"attempt={attempt}/{_TRANSCRIBE_CHUNK_RETRIES}, restart_from={round(start_sec, 2)}s"
                        ),
                    )

                chunk_path = extract_audio_chunk(input_path, start_sec, duration_sec, suffix)
                transcribe_kwargs = {"language": language, "vad_filter": True}
                if initial_prompt:
                    transcribe_kwargs["initial_prompt"] = initial_prompt
                segments, info = whisper_model.transcribe(chunk_path, **transcribe_kwargs)
                if not detected_language:
                    detected_language = info.language

                chunk_segment_count = 0
                last_reported_progress = max(0, int(start_sec / total_duration * 100))
                last_reported_time = start_sec
                for seg in segments:
                    seg_end_overall = min(total_duration, start_sec + float(seg.end))

                    # Interpolate coarse segment boundaries into smoother pseudo progress updates.
                    next_tick = last_reported_time + _TRANSCRIBE_PROGRESS_STEP_SECONDS
                    while next_tick < seg_end_overall:
                        pseudo_progress = min(99, int((next_tick / total_duration) * 100))
                        if pseudo_progress > last_reported_progress and progress_cb:
                            progress_cb(pseudo_progress, chunk_index + 1, total_chunks, next_tick)
                            last_reported_progress = pseudo_progress
                        next_tick += _TRANSCRIBE_PROGRESS_STEP_SECONDS

                    pseudo_progress = min(99, int((seg_end_overall / total_duration) * 100))
                    if pseudo_progress > last_reported_progress and progress_cb:
                        progress_cb(pseudo_progress, chunk_index + 1, total_chunks, seg_end_overall)
                        last_reported_progress = pseudo_progress
                    last_reported_time = max(last_reported_time, seg_end_overall)

                    text = seg.text.strip()
                    if not text:
                        continue

                    abs_start = round(start_sec + float(seg.start), 2)
                    abs_end = round(start_sec + float(seg.end), 2)
                    transcript_parts.append(text)
                    timed_segments.append({"start": abs_start, "end": abs_end, "text": text})
                    chunk_segment_count += 1

                progress = min(99, int(min(total_duration, start_sec + duration_sec) / total_duration * 100))
                if progress_cb and progress > last_reported_progress:
                    progress_cb(progress, chunk_index + 1, total_chunks, start_sec + duration_sec)

                if logger:
                    logger(
                        "debug",
                        (
                            f"transcribe chunk done: chunk={chunk_index + 1}/{total_chunks}, "
                            f"segments={chunk_segment_count}, elapsed_to={round(start_sec + duration_sec, 2)}s"
                        ),
                    )

                last_error = None
                break
            except Exception as exc:
                last_error = exc
                if logger:
                    logger(
                        "error",
                        (
                            f"transcribe chunk failed: chunk={chunk_index + 1}/{total_chunks}, "
                            f"attempt={attempt}/{_TRANSCRIBE_CHUNK_RETRIES}, "
                            f"restart_from={round(start_sec, 2)}s, error={str(exc)}"
                        ),
                    )
                if attempt >= _TRANSCRIBE_CHUNK_RETRIES:
                    raise
            finally:
                if chunk_path:
                    try:
                        os.remove(chunk_path)
                    except OSError:
                        pass

        if last_error is not None:
            raise last_error

    transcript = "\n".join(transcript_parts)
    if logger:
        logger(
            "info",
            (
                f"chunked transcribe complete: chunks={total_chunks}, "
                f"segments={len(timed_segments)}, chars={len(transcript)}"
            ),
        )
    return {
        "language": detected_language or language,
        "duration": round(total_duration, 2) if total_duration else None,
        "text": transcript,
        "segments": timed_segments,
    }


def simple_local_summary(text: str, style: str) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        return "要約対象のテキストがありません。"

    preview = lines[:10]
    if style == "minutes":
        return "\n".join([
            "会議要約",
            "決定事項",
            *[f"{idx + 1}. {line}" for idx, line in enumerate(preview[:4])],
            "論点",
            *[f"{idx + 1}. {line}" for idx, line in enumerate(preview[4:7])],
            "TODO候補",
            *[f"{idx + 1}. {line}" for idx, line in enumerate(preview[7:10])],
        ])
    if style == "actions":
        return "\n".join([
            "実行アクション",
            *[f"{idx + 1}. {line}" for idx, line in enumerate(preview[:8])],
            "補足",
            "担当者と期限は原文から確認してください。",
        ])

    return "\n".join([
        "要点",
        *[f"{idx + 1}. {line}" for idx, line in enumerate(preview[:8])],
    ])


def get_local_summary_model(model_repo: str, gguf_file: str, n_ctx: int, n_threads: int) -> dict:
    global Llama, hf_hub_download
    if Llama is None or hf_hub_download is None:
        try:
            from llama_cpp import Llama as _Llama
            from huggingface_hub import hf_hub_download as _hf_hub_download

            Llama = _Llama
            hf_hub_download = _hf_hub_download
        except Exception as exc:
            raise RuntimeError("ローカルGGUF要約に必要な依存(llama-cpp-python/huggingface_hub)が不足しています。") from exc

    cache_key = f"{model_repo}:{gguf_file}:{n_ctx}:{n_threads}"

    with _local_summary_model_lock:
        cached = _local_summary_model_cache.get(cache_key)
        if cached:
            return cached

        model_cache_dir = MODEL_DIR / "gguf"
        model_cache_dir.mkdir(parents=True, exist_ok=True)

        model_path = hf_hub_download(
            repo_id=model_repo,
            filename=gguf_file,
            local_dir=str(model_cache_dir),
            local_dir_use_symlinks=False,
        )

        llm = Llama(
            model_path=str(model_path),
            n_ctx=max(2048, int(n_ctx)),
            n_threads=max(1, int(n_threads)),
            n_gpu_layers=0,
            verbose=False,
        )

        bundle = {"llm": llm, "model_path": str(model_path)}
        _local_summary_model_cache[cache_key] = bundle
        return bundle


def summarize_with_local_llm(
    system_prompt: str,
    prompt: str,
    model_repo: str,
    gguf_file: str,
    max_new_tokens: int,
    temperature: float,
    context_length: int,
    threads: int,
) -> str:
    bundle = get_local_summary_model(model_repo, gguf_file, context_length, threads)
    llm = bundle["llm"]

    # Some GGUF chat templates reject the system role; merge instructions into user content.
    merged_user_prompt = f"{system_prompt}\n\n{prompt}".strip()
    messages = [
        {"role": "user", "content": merged_user_prompt},
    ]

    completion = llm.create_chat_completion(
        messages=messages,
        max_tokens=max(64, int(max_new_tokens)),
        temperature=max(0.0, float(temperature)),
        top_p=0.9,
    )

    generated_text = (
        completion.get("choices", [{}])[0]
        .get("message", {})
        .get("content", "")
        .strip()
    )
    if not generated_text:
        raise RuntimeError("ローカルLLMの出力が空でした。")
    return generated_text


def normalize_prompt_template(template: str) -> str:
    if not template:
        return DEFAULT_SUMMARY_USER_PROMPT_TEMPLATE
    return template.replace("\\n", "\n")


def extract_text_from_pdf(file_path: str) -> str:
    """PDF ファイルからテキストを抽出"""
    if PdfReader is None:
        raise RuntimeError("pypdf ライブラリが利用できません。")
    
    try:
        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n".join(text_parts)
    except Exception as exc:
        raise RuntimeError(f"PDF テキスト抽出に失敗しました: {str(exc)}")


def extract_text_from_docx(file_path: str) -> str:
    """Word ファイル (.docx) からテキストを抽出"""
    if Document is None:
        raise RuntimeError("python-docx ライブラリが利用できません。")
    
    try:
        doc = Document(file_path)
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        return "\n".join(text_parts)
    except Exception as exc:
        raise RuntimeError(f"Word ファイルのテキスト抽出に失敗しました: {str(exc)}")


def extract_text_from_resume(file_path: str, file_name: str) -> str:
    """レジュメファイルからテキストを抽出（PDF or DOCX に対応）"""
    file_lower = file_name.lower()
    if file_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_path)
    elif file_lower.endswith(".docx"):
        return extract_text_from_docx(file_path)
    else:
        raise RuntimeError(f"対応していないファイル形式です: {file_name}")


def build_summary_prompt(text: str, style: str, prompt_template: str) -> str:
    style_map = {
        "bullets": "重要ポイントを箇条書きで簡潔に整理してください。",
        "minutes": "議事録形式で、決定事項・論点・TODOを分けてください。",
        "actions": "実行アクションを担当・期限つきで抽出してください。",
    }
    style_instruction = style_map.get(style, style_map["bullets"])
    normalized_template = normalize_prompt_template(prompt_template)
    if "{style_instruction}" not in normalized_template:
        normalized_template = "{style_instruction}\n\n" + normalized_template
    if "{text}" not in normalized_template:
        normalized_template = normalized_template + "\n\n{text}"

    try:
        return normalized_template.format(style_instruction=style_instruction, text=text)
    except KeyError:
        return DEFAULT_SUMMARY_USER_PROMPT_TEMPLATE.format(style_instruction=style_instruction, text=text)


def summarize_with_openai(system_prompt: str, prompt: str, model: str, api_key: str) -> str:
    if OpenAI is None:
        raise RuntimeError("openai SDK が利用できません。")

    client = OpenAI(api_key=api_key)
    completion = client.responses.create(
        model=model,
        input=[
            {
                "role": "system",
                "content": system_prompt,
            },
            {
                "role": "user",
                "content": prompt,
            },
        ],
    )
    return completion.output_text.strip()


def summarize_with_gemini(system_prompt: str, prompt: str, model: str, api_key: str) -> str:
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent"
    parts = [
        {
            "text": (
                f"{system_prompt}\n\n"
                f"{prompt}"
            )
        }
    ]

    response = requests.post(
        url,
        params={"key": api_key},
        json={
            "contents": [
                {
                    "parts": parts
                }
            ]
        },
        timeout=60,
    )
    response.raise_for_status()

    payload = response.json()
    candidates = payload.get("candidates", [])
    if not candidates:
        raise RuntimeError("Gemini の応答が空です。")

    parts = candidates[0].get("content", {}).get("parts", [])
    text_parts = [part.get("text", "") for part in parts if part.get("text")]
    if not text_parts:
        raise RuntimeError("Gemini 応答からテキストを取得できませんでした。")
    return "\n".join(text_parts).strip()


def summarize_with_gemini_raw_file(
    system_prompt: str,
    prompt: str,
    model: str,
    api_key: str,
    file_bytes: bytes,
    file_mime_type: str,
) -> str:
    if not file_bytes:
        raise RuntimeError("RAWファイルが空です。")

    # Keep payload under typical inline_data limits in Gemini APIs.
    max_bytes = 10 * 1024 * 1024
    if len(file_bytes) > max_bytes:
        raise RuntimeError("RAWモードのファイルサイズ上限(10MB)を超えています。")

    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent"
    encoded = base64.b64encode(file_bytes).decode("ascii")
    response = requests.post(
        url,
        params={"key": api_key},
        json={
            "contents": [
                {
                    "parts": [
                        {
                            "text": (
                                f"{system_prompt}\n\n"
                                f"以下に添付したレジュメファイルの内容を要約してください。\n"
                                f"{prompt}"
                            )
                        },
                        {
                            "inline_data": {
                                "mime_type": file_mime_type,
                                "data": encoded,
                            }
                        },
                    ]
                }
            ]
        },
        timeout=90,
    )
    response.raise_for_status()

    payload = response.json()
    candidates = payload.get("candidates", [])
    if not candidates:
        raise RuntimeError("Gemini の応答が空です。")

    parts = candidates[0].get("content", {}).get("parts", [])
    text_parts = [part.get("text", "") for part in parts if part.get("text")]
    if not text_parts:
        raise RuntimeError("Gemini 応答からテキストを取得できませんでした。")
    return "\n".join(text_parts).strip()


def detect_resume_mime_type(filename: str) -> str:
    lowered = (filename or "").lower()
    if lowered.endswith(".pdf"):
        return "application/pdf"
    if lowered.endswith(".docx"):
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return "application/octet-stream"


def resolve_summary_provider(provider: str, cfg: dict) -> str:
    requested = (provider or cfg.get("summary_provider") or "auto").lower()
    if requested not in {"auto", "openai", "gemini", "local"}:
        requested = "auto"

    if requested == "auto":
        if cfg.get("gemini_api_key"):
            return "gemini"
        if cfg.get("openai_api_key") and OpenAI is not None:
            return "openai"
        return "local"
    return requested


def llm_summary(
    text: str,
    style: str,
    provider: str,
    model_override: str,
    api_key_override: str,
    system_prompt_override: str,
    user_prompt_template_override: str,
) -> tuple[str, str]:
    cfg = read_config()
    system_prompt = (system_prompt_override or cfg.get("summary_system_prompt") or DEFAULT_SUMMARY_SYSTEM_PROMPT).strip()
    user_prompt_template = (
        user_prompt_template_override
        or cfg.get("summary_user_prompt_template")
        or DEFAULT_SUMMARY_USER_PROMPT_TEMPLATE
    ).strip()
    prompt = build_summary_prompt(text, style, user_prompt_template)
    if "Markdown記法を使わず" not in system_prompt:
        system_prompt = system_prompt + " Markdown記法を使わず、プレーンテキストで出力してください。"
    selected_provider = resolve_summary_provider(provider, cfg)

    if selected_provider == "local":
        local_model = (cfg.get("local_summary_model") or LOCAL_SUMMARY_MODEL).strip()
        local_gguf_file = (cfg.get("local_summary_gguf_file") or LOCAL_SUMMARY_GGUF_FILE).strip()
        local_max_new_tokens = cfg.get("local_summary_max_new_tokens") or LOCAL_SUMMARY_MAX_NEW_TOKENS
        local_temperature = cfg.get("local_summary_temperature")
        local_context_length = cfg.get("local_summary_context_length") or LOCAL_SUMMARY_CONTEXT_LENGTH
        local_threads = cfg.get("local_summary_threads") or LOCAL_SUMMARY_THREADS
        if local_temperature is None:
            local_temperature = LOCAL_SUMMARY_TEMPERATURE
        try:
            summary = summarize_with_local_llm(
                system_prompt=system_prompt,
                prompt=prompt,
                model_repo=local_model,
                gguf_file=local_gguf_file,
                max_new_tokens=int(local_max_new_tokens),
                temperature=float(local_temperature),
                context_length=int(local_context_length),
                threads=int(local_threads),
            )
            return summary, f"local:{local_model}/{local_gguf_file}"
        except Exception as exc:
            add_debug_log("warning", f"local llm summarize failed: {str(exc)}")
            return simple_local_summary(text, style), "local-fallback"

    if selected_provider == "openai":
        model = (model_override or cfg.get("openai_model") or OPENAI_MODEL).strip()
        api_key = (api_key_override or cfg.get("openai_api_key") or "").strip()
        if not api_key:
            return simple_local_summary(text, style), "local-fallback"
        try:
            return summarize_with_openai(system_prompt, prompt, model, api_key), f"openai:{model}"
        except Exception:
            return simple_local_summary(text, style), "local-fallback"

    if selected_provider == "gemini":
        model = (model_override or cfg.get("gemini_model") or GEMINI_MODEL).strip()
        api_key = (api_key_override or cfg.get("gemini_api_key") or "").strip()
        if not api_key:
            return simple_local_summary(text, style), "local-fallback"
        try:
            return summarize_with_gemini(system_prompt, prompt, model, api_key), f"gemini:{model}"
        except Exception:
            return simple_local_summary(text, style), "local-fallback"

    return simple_local_summary(text, style), "local-fallback"


@app.get("/api/health")
def health_check() -> JSONResponse:
    cfg = read_config()
    selected = resolve_summary_provider("", cfg)
    if selected == "openai":
        mode = f"openai:{cfg.get('openai_model') or OPENAI_MODEL}"
    elif selected == "gemini":
        mode = f"gemini:{cfg.get('gemini_model') or GEMINI_MODEL}"
    else:
        mode = (
            "local:"
            f"{cfg.get('local_summary_model') or LOCAL_SUMMARY_MODEL}/"
            f"{cfg.get('local_summary_gguf_file') or LOCAL_SUMMARY_GGUF_FILE}"
        )

    return JSONResponse(
        {
            "status": "ok",
            "whisper_model": cfg.get("whisper_model") or WHISPER_MODEL_NAME,
            "whisper_models": AVAILABLE_WHISPER_MODELS,
            "ffmpeg_available": ffmpeg_available(),
            "summary_mode": mode,
            "summary_provider": cfg.get("summary_provider", "auto"),
        }
    )


@app.get("/api/config")
def get_config() -> JSONResponse:
    return JSONResponse(read_config())


@app.post("/api/config")
def save_config(payload: AppConfigPayload) -> JSONResponse:
    saved = write_config(payload)
    add_debug_log("info", f"config saved: provider={saved.get('summary_provider')}")
    return JSONResponse({"saved": True, "config": saved})


@app.get("/api/debug/logs")
def get_debug_logs() -> JSONResponse:
    return JSONResponse({"logs": _debug_logs})


@app.post("/api/debug/logs/clear")
def clear_debug_logs() -> JSONResponse:
    _debug_logs.clear()
    return JSONResponse({"cleared": True})


@app.post("/api/transcribe")
async def transcribe_audio(
    file: UploadFile = File(...),
    language: str = Form("ja"),
    model: str = Form(""),
) -> JSONResponse:
    if not file.filename:
        raise HTTPException(status_code=400, detail="ファイル名が不正です。")

    if not ffmpeg_available():
        add_debug_log("error", "ffmpeg not found")
        raise HTTPException(status_code=503, detail="FFmpeg が見つかりません。インストール後に再実行してください。")

    cfg = read_config()
    selected_model = (model or cfg.get("whisper_model") or WHISPER_MODEL_NAME).strip().lower()
    if selected_model not in AVAILABLE_WHISPER_MODELS:
        raise HTTPException(
            status_code=400,
            detail=f"Whisperモデルは {', '.join(AVAILABLE_WHISPER_MODELS)} から選択してください。",
        )

    add_debug_log(
        "info",
        f"transcribe start: file={file.filename}, language={language}, model={selected_model}",
    )

    suffix = Path(file.filename).suffix or ".tmp"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        content = await file.read()
        tmp.write(content)
        tmp_path = tmp.name
        add_debug_log("debug", f"temporary file created: {tmp_path}, bytes={len(content)}")

    try:
        add_debug_log("debug", "running chunked whisper transcribe")
        transcribed = transcribe_with_chunk_restart(
            input_path=tmp_path,
            selected_model=selected_model,
            language=language,
            logger=add_debug_log,
        )
        add_debug_log(
            "info",
            (
                f"transcribe done: language={transcribed.get('language', language)}, "
                f"segments={len(transcribed.get('segments', []))}, chars={len(transcribed.get('text', ''))}"
            ),
        )
        return JSONResponse(
            {
                "model": selected_model,
                "language": transcribed.get("language", language),
                "duration": transcribed.get("duration"),
                "text": transcribed.get("text", ""),
                "segments": transcribed.get("segments", []),
            }
        )
    except Exception as exc:
        add_debug_log("error", f"transcribe failed: {str(exc)}")
        raise HTTPException(status_code=500, detail=f"文字起こしに失敗しました: {str(exc)}")
    finally:
        try:
            os.remove(tmp_path)
            add_debug_log("debug", f"temporary file removed: {tmp_path}")
        except OSError:
            pass


@app.post("/api/transcribe/start")
async def start_transcribe_job(
    file: UploadFile = File(...),
    language: str = Form("ja"),
    model: str = Form(""),
    initial_prompt: str = Form(""),
) -> JSONResponse:
    if not file.filename:
        raise HTTPException(status_code=400, detail="ファイル名が不正です。")

    if not ffmpeg_available():
        add_debug_log("error", "ffmpeg not found")
        raise HTTPException(status_code=503, detail="FFmpeg が見つかりません。インストール後に再実行してください。")

    cfg = read_config()
    selected_model = (model or cfg.get("whisper_model") or WHISPER_MODEL_NAME).strip().lower()
    if selected_model not in AVAILABLE_WHISPER_MODELS:
        raise HTTPException(
            status_code=400,
            detail=f"Whisperモデルは {', '.join(AVAILABLE_WHISPER_MODELS)} から選択してください。",
        )

    suffix = Path(file.filename).suffix or ".tmp"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        content = await file.read()
        tmp.write(content)
        tmp_path = tmp.name

    initial_prompt_normalized = (initial_prompt or "").strip()
    job_id = create_transcribe_job(
        selected_model=selected_model,
        language=language,
        filename=file.filename,
        initial_prompt=initial_prompt_normalized,
    )
    add_debug_log(
        "info",
        f"transcribe start: job={job_id}, file={file.filename}, language={language}, model={selected_model}, has_prompt={bool(initial_prompt_normalized)}",
    )
    add_debug_log("debug", f"temporary file created: {tmp_path}, bytes={len(content)}")

    thread = threading.Thread(
        target=run_transcribe_job,
        args=(job_id, tmp_path, selected_model, language, initial_prompt_normalized),
        daemon=True,
    )
    thread.start()

    return JSONResponse(
        {
            "job_id": job_id,
            "status": "queued",
            "progress": 0,
            "message": "文字起こしジョブを開始しました。",
        }
    )


@app.get("/api/transcribe/jobs/{job_id}")
def get_transcribe_job_status(job_id: str) -> JSONResponse:
    job = get_transcribe_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="ジョブが見つかりません。")

    response = {
        "job_id": job["job_id"],
        "status": job["status"],
        "progress": job.get("progress", 0),
        "message": job.get("message", ""),
        "model": job.get("model"),
        "language": job.get("language"),
    }

    if job["status"] == "completed":
        response["result"] = job.get("result")
    if job["status"] == "failed":
        response["error"] = job.get("error") or "unknown error"

    return JSONResponse(response)


@app.post("/api/extract-resume")
async def extract_resume(file: UploadFile = File(...)) -> JSONResponse:
    """レジュメファイル（PDF/DOCX）からテキストを抽出"""
    if not file.filename:
        raise HTTPException(status_code=400, detail="ファイルを選択してください。")
    
    # ファイル形式チェック
    file_lower = file.filename.lower()
    if not (file_lower.endswith(".pdf") or file_lower.endswith(".docx")):
        raise HTTPException(
            status_code=400, 
            detail="PDF または Word ファイル (.docx) のみ対応しています。"
        )
    
    # 一時ファイルに保存
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=file.filename) as tmp_file:
            tmp_path = tmp_file.name
            content = await file.read()
            tmp_file.write(content)
        
        # テキスト抽出
        extracted_text = extract_text_from_resume(tmp_path, file.filename)
        if not extracted_text.strip():
            raise HTTPException(
                status_code=400,
                detail="ファイルからテキストを抽出できませんでした。"
            )
        
        add_debug_log("info", f"resume extracted: filename={file.filename}, chars={len(extracted_text)}")
        return JSONResponse({
            "text": extracted_text,
            "filename": file.filename
        })
    except HTTPException:
        raise
    except Exception as exc:
        add_debug_log("error", f"resume extraction failed: filename={file.filename}, error={str(exc)}")
        raise HTTPException(status_code=500, detail=f"ファイル処理エラー: {str(exc)}")
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except OSError:
                pass


@app.post("/api/summarize")
async def summarize_text(
    text: str = Form(""),
    style: str = Form("bullets"),
    provider: str = Form("auto"),
    resume_mode: str = Form("extract"),
    resume_file: UploadFile | None = File(None),
    model: str = Form(""),
    api_key: str = Form(""),
    system_prompt: str = Form(""),
    user_prompt_template: str = Form(""),
) -> JSONResponse:
    normalized_resume_mode = (resume_mode or "extract").strip().lower()
    has_resume_file = resume_file is not None and bool(resume_file.filename)

    # Backward-compatible fallback:
    # If a file is attached and text is empty, treat it as RAW mode even when
    # resume_mode was not sent/selected correctly on the client.
    if has_resume_file and not (text or "").strip():
        normalized_resume_mode = "raw"

    if normalized_resume_mode not in {"extract", "raw"}:
        add_debug_log("warning", f"summarize invalid resume_mode: {resume_mode}")
        raise HTTPException(status_code=400, detail="resume_mode は extract または raw を指定してください。")

    if normalized_resume_mode == "raw":
        if not has_resume_file:
            add_debug_log("warning", "summarize raw rejected: resume_file missing")
            raise HTTPException(status_code=400, detail="RAWモードではレジュメファイルが必須です。")
        if not text.strip():
            text = "添付ファイルの内容を要約してください。"
    elif not text.strip():
        raise HTTPException(status_code=400, detail="要約対象のテキストが空です。")

    add_debug_log(
        "info",
        f"summarize start: provider={provider}, style={style}, chars={len(text)}, resume_mode={normalized_resume_mode}",
    )

    if normalized_resume_mode == "raw":
        cfg = read_config()
        selected_provider = resolve_summary_provider(provider, cfg)
        if selected_provider != "gemini":
            add_debug_log("warning", f"summarize raw rejected: provider={selected_provider}")
            raise HTTPException(
                status_code=400,
                detail="RAWモードは現在 Gemini プロバイダのみ対応しています。",
            )

        model_name = (model or cfg.get("gemini_model") or GEMINI_MODEL).strip()
        api_key_value = (api_key or cfg.get("gemini_api_key") or "").strip()
        if not api_key_value:
            add_debug_log("warning", "summarize raw rejected: gemini api key missing")
            raise HTTPException(status_code=400, detail="Gemini APIキーが未設定です。")

        system_prompt_value = (system_prompt or cfg.get("summary_system_prompt") or DEFAULT_SUMMARY_SYSTEM_PROMPT).strip()
        user_prompt_template_value = (
            user_prompt_template
            or cfg.get("summary_user_prompt_template")
            or DEFAULT_SUMMARY_USER_PROMPT_TEMPLATE
        ).strip()
        prompt = build_summary_prompt(text, style, user_prompt_template_value)

        file_bytes = await resume_file.read()
        file_mime = detect_resume_mime_type(resume_file.filename)
        try:
            summary = summarize_with_gemini_raw_file(
                system_prompt=system_prompt_value,
                prompt=prompt,
                model=model_name,
                api_key=api_key_value,
                file_bytes=file_bytes,
                file_mime_type=file_mime,
            )
            add_debug_log(
                "info",
                (
                    "summarize done: "
                    f"mode=gemini:{model_name}+raw, chars={len(summary)}, "
                    f"resume_file={resume_file.filename}, bytes={len(file_bytes)}"
                ),
            )
            return JSONResponse({
                "summary": summary,
                "style": style,
                "summary_mode": f"gemini:{model_name}+raw",
            })
        except Exception as exc:
            add_debug_log("error", f"summarize raw failed: error={str(exc)}")
            raise HTTPException(status_code=500, detail=f"RAW要約に失敗しました: {str(exc)}")

    summary, summary_mode = llm_summary(
        text,
        style,
        provider,
        model,
        api_key,
        system_prompt,
        user_prompt_template,
    )
    add_debug_log("info", f"summarize done: mode={summary_mode}, chars={len(summary)}")
    return JSONResponse({"summary": summary, "style": style, "summary_mode": summary_mode})


if WEB_DIR.exists():
    app.mount("/", StaticFiles(directory=str(WEB_DIR), html=True), name="web")
